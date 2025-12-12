"""
Test suite for the Word to Text processor.

Tests:
1. Basic .docx to text conversion
2. Handling of tables in Word documents
3. Preservation of formatting
4. Error handling for missing dependencies
5. Metadata generation
"""

import pytest
import os
import asyncio
from pathlib import Path
from docex.docbasket import DocBasket
from docex.processors.word_to_text import WordToTextProcessor
from docex.processors.base import ProcessingResult
from datetime import datetime


def create_test_word_document(file_path: str) -> Path:
    """Create a test Word document for testing"""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = DocxDocument()

    # Add title
    doc.add_heading('Test Document', level=1)

    # Add paragraphs
    doc.add_paragraph('This is the first paragraph.')
    doc.add_paragraph('This is the second paragraph with some content.')
    doc.add_paragraph('')  # Empty paragraph for spacing

    # Add a section header
    doc.add_heading('Section with Table', level=2)

    # Add a table
    table = doc.add_table(rows=3, cols=3)
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Column 1'
    header_cells[1].text = 'Column 2'
    header_cells[2].text = 'Column 3'

    # Add data rows
    row1_cells = table.rows[1].cells
    row1_cells[0].text = 'Data 1A'
    row1_cells[1].text = 'Data 1B'
    row1_cells[2].text = 'Data 1C'

    row2_cells = table.rows[2].cells
    row2_cells[0].text = 'Data 2A'
    row2_cells[1].text = 'Data 2B'
    row2_cells[2].text = 'Data 2C'

    # Save document
    doc.save(file_path)
    return Path(file_path)


def test_basic_word_to_text_conversion():
    """Test basic conversion of a Word document to text"""
    print("\n=== Testing Basic Word to Text Conversion ===\n")

    # Create test basket
    basket_name = f"test_basket_word_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    print(f"Creating test basket: {basket_name}")
    test_basket = DocBasket.create(basket_name)

    # Create test Word document
    test_file = "test_document.docx"
    print(f"Creating test Word document: {test_file}")
    create_test_word_document(test_file)

    try:
        # Add document to basket
        print("Adding document to basket...")
        doc = test_basket.add(test_file)
        print(f"Document added: {doc.name}")

        # Create processor with default config
        processor = WordToTextProcessor(config={
            'encoding': 'utf-8',
            'preserve_formatting': True,
            'extract_tables': True
        })

        # Verify processor can handle the document
        assert processor.can_process(doc), "Processor should be able to process .docx files"

        # Process document
        print("Processing document...")
        result = asyncio.run(processor.process(doc))

        # Verify processing was successful
        assert result.success, f"Processing should succeed: {result.error}"
        assert result.content is not None, "Result should have content"
        assert result.metadata is not None, "Result should have metadata"

        print(f"✓ Processing successful!")
        print(f"Output file: {result.content}")

        # Verify metadata
        print("\nVerifying metadata...")
        assert result.metadata['input_format'] == 'docx'
        assert result.metadata['output_format'] == 'txt'
        assert result.metadata['paragraph_count'] > 0
        assert result.metadata['word_count'] > 0
        assert result.metadata['char_count'] > 0

        print(f"  Paragraphs: {result.metadata['paragraph_count']}")
        print(f"  Tables: {result.metadata['table_count']}")
        print(f"  Words: {result.metadata['word_count']}")
        print(f"  Characters: {result.metadata['char_count']}")

        # Verify output file exists
        # The result.content is the storage path, construct the full filesystem path
        from pathlib import Path as PathLib
        storage_base = PathLib("storage")
        output_path = storage_base / result.content
        assert output_path.exists(), f"Output file should exist at: {output_path}"

        # Read and verify content
        with open(output_path, 'r', encoding='utf-8') as f:
            text_content = f.read()

        print("\n" + "="*60)
        print("Extracted Text:")
        print("="*60)
        print(text_content)
        print("="*60)

        assert 'Test Document' in text_content
        assert 'first paragraph' in text_content
        assert 'second paragraph' in text_content
        assert 'Table 1' in text_content  # Table marker
        assert 'Column 1' in text_content  # Table header
        assert 'Data 1A' in text_content  # Table data

        print("\n✓ All assertions passed!")

    finally:
        # Clean up test file
        if os.path.exists(test_file):
            os.remove(test_file)
            print(f"\nCleaned up test file: {test_file}")


def test_word_processor_can_process():
    """Test the can_process method"""
    print("\n=== Testing can_process Method ===\n")

    # Create test basket
    basket_name = f"test_basket_can_process_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    test_basket = DocBasket.create(basket_name)

    processor = WordToTextProcessor(config={})

    # Test .docx file
    docx_file = "test.docx"
    with open(docx_file, 'w') as f:
        f.write("dummy")

    try:
        doc_docx = test_basket.add(docx_file)
        assert processor.can_process(doc_docx), "Should process .docx files"
        print("✓ Correctly identifies .docx files")
    finally:
        if os.path.exists(docx_file):
            os.remove(docx_file)

    # Test .doc file
    doc_file = "test.doc"
    with open(doc_file, 'w') as f:
        f.write("dummy")

    try:
        doc_doc = test_basket.add(doc_file)
        assert processor.can_process(doc_doc), "Should process .doc files"
        print("✓ Correctly identifies .doc files")
    finally:
        if os.path.exists(doc_file):
            os.remove(doc_file)

    # Test non-Word file
    txt_file = "test.txt"
    with open(txt_file, 'w') as f:
        f.write("dummy")

    try:
        doc_txt = test_basket.add(txt_file)
        assert not processor.can_process(doc_txt), "Should not process .txt files"
        print("✓ Correctly rejects non-Word files")
    finally:
        if os.path.exists(txt_file):
            os.remove(txt_file)


def test_word_to_text_no_tables():
    """Test conversion with table extraction disabled"""
    print("\n=== Testing Word to Text without Table Extraction ===\n")

    basket_name = f"test_basket_no_tables_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    test_basket = DocBasket.create(basket_name)

    test_file = "test_no_tables.docx"
    create_test_word_document(test_file)

    try:
        doc = test_basket.add(test_file)

        # Create processor with tables disabled
        processor = WordToTextProcessor(config={
            'encoding': 'utf-8',
            'preserve_formatting': True,
            'extract_tables': False  # Disable table extraction
        })

        result = asyncio.run(processor.process(doc))

        assert result.success, "Processing should succeed"

        # Read output
        from pathlib import Path as PathLib
        storage_base = PathLib("storage")
        output_path = storage_base / result.content

        with open(output_path, 'r', encoding='utf-8') as f:
            text_content = f.read()

        # Should not have table markers
        assert '--- Table' not in text_content
        print("✓ Tables correctly excluded from output")

    finally:
        if os.path.exists(test_file):
            os.remove(test_file)


def test_word_to_text_no_formatting():
    """Test conversion with formatting preservation disabled"""
    print("\n=== Testing Word to Text without Formatting Preservation ===\n")

    basket_name = f"test_basket_no_format_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    test_basket = DocBasket.create(basket_name)

    test_file = "test_no_format.docx"
    create_test_word_document(test_file)

    try:
        doc = test_basket.add(test_file)

        # Create processor with formatting disabled
        processor = WordToTextProcessor(config={
            'encoding': 'utf-8',
            'preserve_formatting': False,  # Disable formatting
            'extract_tables': True
        })

        result = asyncio.run(processor.process(doc))

        assert result.success, "Processing should succeed"
        print("✓ Processing without formatting preservation succeeded")

    finally:
        if os.path.exists(test_file):
            os.remove(test_file)


def test_word_processor_error_handling():
    """Test error handling for legacy .doc files"""
    print("\n=== Testing Error Handling for .doc Files ===\n")

    basket_name = f"test_basket_error_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    test_basket = DocBasket.create(basket_name)

    # Create a fake .doc file (not really a .doc, just for testing)
    test_file = "test_legacy.doc"
    with open(test_file, 'w') as f:
        f.write("This is not a real .doc file")

    try:
        doc = test_basket.add(test_file)

        processor = WordToTextProcessor(config={})

        result = asyncio.run(processor.process(doc))

        # Should fail because .doc is not supported
        assert not result.success, "Processing .doc should fail"
        assert "Legacy .doc format is not supported" in result.error
        print("✓ Correctly handles unsupported .doc format")

    finally:
        if os.path.exists(test_file):
            os.remove(test_file)


if __name__ == "__main__":
    """Run all tests"""
    print("\n" + "="*70)
    print("Word to Text Processor Test Suite")
    print("="*70)

    try:
        test_basic_word_to_text_conversion()
        test_word_processor_can_process()
        test_word_to_text_no_tables()
        test_word_to_text_no_formatting()
        test_word_processor_error_handling()

        print("\n" + "="*70)
        print("✓ ALL TESTS PASSED!")
        print("="*70)

    except Exception as e:
        print("\n" + "="*70)
        print("✗ TEST FAILED!")
        print("="*70)
        print(f"Error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise
