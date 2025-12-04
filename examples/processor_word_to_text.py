"""
Example demonstrating the use of the Word to Text processor.

This example:
1. Creates a test .docx file
2. Creates a document from the Word file
3. Processes the document using the Word to Text processor
4. Displays the processing results
5. Uses UserContext for audit logging

Security Best Practices:
- Always use UserContext for audit logging
- UserContext enables operation tracking

Requirements:
- python-docx library (install with: pip install python-docx)
"""

import os
from pathlib import Path
from docex import DocEX
from docex.context import UserContext
from docex.processors.factory import factory
from docex.processors.word_to_text import WordToTextProcessor
import shutil


def create_test_docx(file_path: Path) -> Path:
    """Create a test Word document with sample content and return the file path"""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
    except ImportError:
        raise ImportError(
            "python-docx library is required. Install it with: pip install python-docx"
        )

    docx_file = file_path / 'test_document.docx'

    # Create a new Document
    doc = DocxDocument()

    # Add title
    title = doc.add_heading('Sample Document', level=1)

    # Add introduction paragraph
    intro = doc.add_paragraph(
        'This is a sample Word document created for testing the Word to Text processor. '
        'It contains multiple paragraphs, formatted text, and a table.'
    )

    # Add section header
    doc.add_heading('Section 1: Overview', level=2)

    # Add content paragraphs
    doc.add_paragraph(
        'The DocEX Word to Text processor can extract text content from Microsoft Word '
        'documents (.docx format) and convert them to plain text files.'
    )

    doc.add_paragraph(
        'Key features include:\n'
        '• Preservation of paragraph structure\n'
        '• Table extraction\n'
        '• Configurable text encoding\n'
        '• Detailed metadata about the conversion process'
    )

    # Add another section
    doc.add_heading('Section 2: Sample Data', level=2)

    # Add a table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Name'
    header_cells[1].text = 'Role'
    header_cells[2].text = 'Department'

    # Data rows
    data = [
        ('John Doe', 'Engineer', 'Engineering'),
        ('Jane Smith', 'Manager', 'Operations'),
        ('Bob Johnson', 'Analyst', 'Finance')
    ]

    for i, (name, role, dept) in enumerate(data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = name
        row_cells[1].text = role
        row_cells[2].text = dept

    # Add conclusion
    doc.add_heading('Conclusion', level=2)
    doc.add_paragraph(
        'This document demonstrates the capabilities of the Word to Text processor '
        'in extracting content from structured Word documents.'
    )

    # Save the document
    doc.save(str(docx_file))
    return docx_file


def main():
    """Run the example"""
    # Create UserContext for audit logging
    user_context = UserContext(
        user_id="word_processor",
        user_email="processor@example.com",
        tenant_id="example_tenant",  # Optional: for multi-tenant applications
        roles=["user"]
    )

    # Initialize DocEX with UserContext (enables audit logging)
    docEX = DocEX(user_context=user_context)

    # Create test directory
    test_dir = Path("test_data")
    test_dir.mkdir(exist_ok=True)

    try:
        # Create test Word file
        print("Creating test Word document...")
        docx_file = create_test_docx(test_dir)
        print(f"Created test Word file: {docx_file}")

        # Get or create basket (same pattern as basic_usage.py)
        basket_name = "word_test_basket"
        try:
            basket = docEX.create_basket(basket_name)
            print(f"Created new basket: {basket_name}")
        except ValueError:
            print(f"Using existing basket: {basket_name}")
            baskets = docEX.list_baskets()
            basket = next((b for b in baskets if b.name == basket_name), None)
            if not basket:
                raise RuntimeError(f"Failed to get basket: {basket_name}")

        # Add document to basket
        print("Adding document to basket...")
        document = basket.add(str(docx_file))
        print(f"Added document to basket: {document.name}")

        # Register the processor class in the factory
        factory.register(WordToTextProcessor)

        # Create processor instance with configuration
        # Configuration options:
        # - encoding: Output text encoding (default: 'utf-8')
        # - preserve_formatting: Keep line breaks and spacing (default: True)
        # - extract_tables: Include table content in output (default: True)
        processor = WordToTextProcessor(config={
            'encoding': 'utf-8',
            'preserve_formatting': True,
            'extract_tables': True
        })

        # Process document
        print("\nProcessing document...")
        import asyncio
        result = asyncio.run(processor.process(document))

        if result.success:
            print("\n✓ Processing successful!")
            print(f"Output file: {result.content}")
            print("\nMetadata:")
            for key, value in result.metadata.items():
                print(f"  {key}: {value}")

            # Read and display the extracted text
            output_path = Path(result.content)
            if output_path.exists():
                print("\n" + "="*60)
                print("EXTRACTED TEXT CONTENT:")
                print("="*60)
                with open(output_path, 'r', encoding='utf-8') as f:
                    print(f.read())
                print("="*60)
        else:
            print("\n✗ Processing failed:", result.error)

    except Exception as e:
        print(f"Error: {str(e)}")
        import traceback
        traceback.print_exc()

    finally:
        # Clean up test files and directories
        if test_dir.exists():
            shutil.rmtree(test_dir)
            print("\nCleaned up test files")


if __name__ == '__main__':
    main()
