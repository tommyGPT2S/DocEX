"""
Word to Text Processor

Converts Microsoft Word documents (.docx and .doc) to plain text files.
Uses python-docx for .docx files and provides fallback for .doc files.
"""

import io
from typing import Dict, Any, Optional
from pathlib import Path
from docex.processors.base import BaseProcessor, ProcessingResult
from docex.document import Document
from docex.db.connection import Database


class WordToTextProcessor(BaseProcessor):
    """Processor that converts Word documents to plain text format"""

    def __init__(self, config: Dict[str, Any], db: Optional[Database] = None):
        """
        Initialize the Word to Text processor

        Args:
            config: Configuration dictionary with optional keys:
                - encoding: Output text encoding (default: 'utf-8')
                - preserve_formatting: Whether to preserve basic formatting like line breaks (default: True)
                - extract_tables: Whether to extract table content (default: True)
            db: Optional database instance for multi-tenancy support
        """
        super().__init__(config, db=db)
        self.encoding = config.get('encoding', 'utf-8')
        self.preserve_formatting = config.get('preserve_formatting', True)
        self.extract_tables = config.get('extract_tables', True)

    def can_process(self, document: Document) -> bool:
        """
        Check if this processor can handle the document

        Args:
            document: Document to check

        Returns:
            True if document is a Word file (.doc or .docx)
        """
        doc_name_lower = document.name.lower()
        return doc_name_lower.endswith('.docx') or doc_name_lower.endswith('.doc')

    async def process(self, document: Document) -> ProcessingResult:
        """
        Process the Word document and convert it to plain text

        Args:
            document: Document to process

        Returns:
            ProcessingResult with success status, output path, and metadata
        """
        try:
            # Import python-docx (fail gracefully if not installed)
            try:
                from docx import Document as DocxDocument
            except ImportError:
                return ProcessingResult(
                    success=False,
                    error="python-docx library is not installed. Install it with: pip install python-docx"
                )

            # Get document content as bytes
            doc_bytes = self.get_document_bytes(document)

            # Check file type
            is_docx = document.name.lower().endswith('.docx')
            is_doc = document.name.lower().endswith('.doc')

            if is_doc and not is_docx:
                # Legacy .doc format - python-docx doesn't support this natively
                return ProcessingResult(
                    success=False,
                    error="Legacy .doc format is not supported. Please convert to .docx first or use an alternative converter."
                )

            # Parse the .docx file
            docx_doc = DocxDocument(io.BytesIO(doc_bytes))

            # Extract text content
            text_parts = []
            paragraph_count = 0
            table_count = 0

            # Extract paragraphs
            for paragraph in docx_doc.paragraphs:
                paragraph_text = paragraph.text
                if paragraph_text.strip():  # Only add non-empty paragraphs
                    text_parts.append(paragraph_text)
                    paragraph_count += 1
                elif self.preserve_formatting:
                    # Preserve empty lines for formatting
                    text_parts.append('')

            # Extract tables if configured
            if self.extract_tables:
                for table in docx_doc.tables:
                    table_count += 1
                    if self.preserve_formatting:
                        text_parts.append('')  # Add spacing before table
                        text_parts.append(f'--- Table {table_count} ---')

                    for row in table.rows:
                        row_texts = [cell.text.strip() for cell in row.cells]
                        row_text = '\t'.join(row_texts)  # Tab-separated values
                        text_parts.append(row_text)

                    if self.preserve_formatting:
                        text_parts.append(f'--- End Table {table_count} ---')
                        text_parts.append('')  # Add spacing after table

            # Join all text parts
            if self.preserve_formatting:
                full_text = '\n'.join(text_parts)
            else:
                # Join with single newline, no extra spacing
                full_text = '\n'.join(part for part in text_parts if part)

            # Create output file name
            # document.path is like: docex/basket_xxx/doc_yyy.docx
            # We need to change it to: basket_xxx/doc_yyy.txt
            # The storage.save() will add the base path (storage/docex) automatically

            doc_path = Path(document.path)
            doc_path_parts = list(doc_path.parts)

            # Remove the 'docex' prefix if present (storage base)
            if doc_path_parts and doc_path_parts[0] == 'docex':
                doc_path_parts = doc_path_parts[1:]

            # Change the extension to .txt
            if doc_path_parts:
                doc_path_parts[-1] = Path(doc_path_parts[-1]).with_suffix('.txt').name

            # Reconstruct the path
            output_path = '/'.join(doc_path_parts)

            # Save the text content
            text_bytes = full_text.encode(self.encoding)
            document.storage_service.storage.save(
                output_path,
                io.BytesIO(text_bytes)
            )

            # Calculate statistics
            word_count = len(full_text.split())
            char_count = len(full_text)
            line_count = len(text_parts)

            # Record the operation in the database
            self._record_operation(
                document=document,
                status='success',
                input_metadata={
                    'input_format': 'docx' if is_docx else 'doc',
                    'input_file': document.name,
                    'input_size': len(doc_bytes)
                },
                output_metadata={
                    'output_format': 'txt',
                    'output_file': str(output_path),
                    'paragraph_count': paragraph_count,
                    'table_count': table_count,
                    'word_count': word_count,
                    'char_count': char_count,
                    'line_count': line_count
                }
            )

            # Create processing result
            # Return the full storage path (including 'docex/' prefix) for consistency
            full_output_path = f"docex/{output_path}"

            return ProcessingResult(
                success=True,
                content=full_output_path,
                metadata={
                    'input_format': 'docx' if is_docx else 'doc',
                    'output_format': 'txt',
                    'output_path': full_output_path,
                    'paragraph_count': paragraph_count,
                    'table_count': table_count,
                    'word_count': word_count,
                    'char_count': char_count,
                    'line_count': line_count,
                    'encoding': self.encoding
                }
            )

        except Exception as e:
            # Record failed operation
            self._record_operation(
                document=document,
                status='failed',
                error=str(e)
            )

            return ProcessingResult(
                success=False,
                error=f"Failed to convert Word document to text: {str(e)}"
            )
